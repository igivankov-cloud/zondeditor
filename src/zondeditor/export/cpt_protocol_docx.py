from __future__ import annotations

import datetime as dt
import subprocess
from pathlib import Path
from typing import Any

try:
    from docx import Document
except Exception:  # pragma: no cover
    Document = None


def _git_hash() -> str:
    try:
        return subprocess.check_output(["git", "rev-parse", "--short", "HEAD"], text=True).strip()
    except Exception:
        return "unknown"


def export_cpt_protocol_docx(*, out_path: Path, object_name: str, settings: dict[str, Any], rows: list[dict[str, Any]], template_path: Path | None = None) -> Path:
    if Document is None:
        raise RuntimeError("python-docx не установлен")
    doc = Document(str(template_path)) if template_path and template_path.exists() else Document()

    doc.add_heading("Протокол CPT: φ и E по СП 446 Прил. Ж", level=1)
    doc.add_paragraph(f"Объект: {object_name or '-'}")
    doc.add_paragraph(f"Дата: {dt.datetime.now().strftime('%d.%m.%Y %H:%M')}")
    doc.add_paragraph(f"Версия программы (git hash): {_git_hash()}")
    method = str(settings.get("method") or "SP446_APP_J")
    method_label = "СП 446.1325800.2019 (Приложение Ж)" if method == "SP446_APP_J" else "СП 11-105-97 (Приложение И)"
    doc.add_paragraph(f"Выбранная методика: {method_label}")
    doc.add_paragraph("Норматив: СП 446.1325800.2019 (ред. по данным установленной БД НД), Приложение Ж.")
    gwl = settings.get("groundwater_level")
    doc.add_paragraph(f"УГВ: {'не задан' if gwl in (None, '') else f'{gwl} м'}")

    for row in rows:
        doc.add_heading(str(row.get("ige_id") or "ИГЭ"), level=2)
        doc.add_paragraph(f"Границы слоя: {row.get('bounds') or '-'}; mid_depth={row.get('mid_depth', '-')} м")
        doc.add_paragraph(f"Тип грунта: {row.get('soil_type') or '-'}")
        doc.add_paragraph(
            f"Признаки: sand_class={row.get('sand_class') or '-'}, alluvial={'да' if row.get('alluvial') else 'нет'}, "
            f"saturated={row.get('saturated') if row.get('saturated') is not None else 'auto/не задан'}, "
            f"IL={row.get('il') or '-'}, консистенция={row.get('consistency') or '-'}"
        )
        sf = dict(row.get("source_flags") or {})
        doc.add_paragraph(f"Источник/флаги: CPT={bool(sf.get('CPT', True))}, LAB={bool(sf.get('LAB', False))}, Stamp={bool(sf.get('Stamp', False))}")
        doc.add_paragraph(
            f"Статистика qc: mean={row.get('qc_mean', '-')}, n={row.get('n', '-')}, min/max={row.get('qc_min', '-')}/{row.get('qc_max', '-')}, "
            f"std={row.get('std', '-')}, V={row.get('variation', '-')}"
        )
        doc.add_paragraph("Проверки ГОСТ 20522-2012: рекомендуется n≥6 и V≤0.30 (информативная проверка).")
        doc.add_paragraph(
            f"Lookup: таблица {row.get('lookup_table', '-')}, ветка: {row.get('lookup_branch', '-')}, диапазон qc: {row.get('lookup_interval', '-')}"
        )
        if row.get("status") == "не рассчитано" or row.get("reason"):
            doc.add_paragraph(f"Итог: не рассчитывается. Причина: {row.get('reason') or 'нет нормирования'}")
        else:
            doc.add_paragraph(f"Итог: φ_norm={row.get('phi_norm', '-')}, E_norm={row.get('E_norm', '-')}")
        if row.get("note"):
            doc.add_paragraph(f"Обоснование/note: {row.get('note')}")

    doc.add_heading("Сводная таблица по ИГЭ", level=2)
    table = doc.add_table(rows=max(1, len(rows)) + 1, cols=6)
    hdr = table.rows[0].cells
    hdr[0].text = "ИГЭ"
    hdr[1].text = "qc_ср"
    hdr[2].text = "φ_norm"
    hdr[3].text = "E_norm"
    hdr[4].text = "Источник"
    hdr[5].text = "Статус"
    for i, row in enumerate(rows, start=1):
        cells = table.rows[i].cells
        sf = dict(row.get("source_flags") or {})
        src = ",".join([k for k, v in {"CPT": sf.get("CPT", True), "LAB": sf.get("LAB", False), "Stamp": sf.get("Stamp", False)}.items() if bool(v)])
        cells[0].text = str(row.get("ige_id") or "")
        cells[1].text = str(row.get("qc_mean") or "")
        cells[2].text = str(row.get("phi_norm") if row.get("phi_norm") is not None else "-")
        cells[3].text = str(row.get("E_norm") if row.get("E_norm") is not None else "-")
        cells[4].text = src or "-"
        cells[5].text = str(row.get("status") or "-")

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path
